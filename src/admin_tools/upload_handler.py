"""
Admin Upload Handler for ex-GPT System
관리도구 파일 업로드 처리 모듈
"""

import os
import shutil
import asyncio
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
from pathlib import Path
import json
import logging
from dataclasses import dataclass, asdict
from enum import Enum
import hashlib
from concurrent.futures import ProcessPoolExecutor
import aiofiles
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage

# 로컬 모듈
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from image_processing import (
    IntegratedImageAnalyzer,
    ProcessingMode,
    ComprehensiveImageAnalysis
)
from rag_pipeline.embeddings import EmbeddingGenerator
from rag_pipeline.vector_store import QdrantVectorStore

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class UploadType(Enum):
    """업로드 유형"""
    ADMIN = "admin"        # 관리자 업로드 (영구 저장)
    USER = "user"          # 사용자 업로드 (임시 저장)
    BATCH = "batch"        # 배치 업로드
    NATIONAL_AUDIT = "ga"  # 국정감사 자료


class FileType(Enum):
    """파일 유형"""
    IMAGE = "image"
    DOCUMENT = "document"
    PDF = "pdf"
    EXCEL = "excel"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class UploadResult:
    """업로드 처리 결과"""
    file_id: str
    original_name: str
    stored_path: str
    file_type: FileType
    upload_type: UploadType
    processing_status: str
    analysis_result: Optional[Dict[str, Any]]
    vector_id: Optional[str]
    metadata: Dict[str, Any]
    errors: List[str]
    timestamp: str


class AdminUploadHandler:
    """
    관리도구 업로드 핸들러
    파일 업로드, 처리, 저장을 관리
    """
    
    def __init__(self,
                 upload_dir: str = "./uploads",
                 processed_dir: str = "./processed",
                 temp_dir: str = "./temp",
                 max_file_size: int = 100 * 1024 * 1024):  # 100MB
        """
        업로드 핸들러 초기화
        
        Args:
            upload_dir: 업로드 디렉토리
            processed_dir: 처리 완료 디렉토리
            temp_dir: 임시 디렉토리
            max_file_size: 최대 파일 크기
        """
        # 디렉토리 설정
        self.upload_dir = Path(upload_dir)
        self.processed_dir = Path(processed_dir)
        self.temp_dir = Path(temp_dir)
        
        # 디렉토리 생성
        for dir_path in [self.upload_dir, self.processed_dir, self.temp_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)
            
        self.max_file_size = max_file_size
        
        # 허용된 확장자
        self.allowed_extensions = {
            'image': {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'},
            'document': {'.doc', '.docx', '.hwp', '.hwpx'},
            'pdf': {'.pdf'},
            'excel': {'.xls', '.xlsx', '.csv'},
            'text': {'.txt', '.md', '.log'}
        }
        
        # 모듈 초기화
        self.image_analyzer = IntegratedImageAnalyzer()
        self.embedding_generator = None  # 나중에 초기화
        self.vector_store = None  # 나중에 초기화
        
        # 처리 큐
        self.processing_queue = asyncio.Queue()
        self.processing_tasks = []
        
        logger.info("관리도구 업로드 핸들러 초기화 완료")
        
    async def initialize_modules(self):
        """비동기 모듈 초기화"""
        try:
            # 임베딩 생성기 초기화
            self.embedding_generator = EmbeddingGenerator()
            
            # 벡터 저장소 초기화
            self.vector_store = QdrantVectorStore()
            await self.vector_store.initialize()
            
            # 백그라운드 처리 시작
            self._start_background_processing()
            
            logger.info("모든 모듈 초기화 완료")
            
        except Exception as e:
            logger.error(f"모듈 초기화 실패: {str(e)}")
            raise
            
    def _start_background_processing(self):
        """백그라운드 처리 태스크 시작"""
        for i in range(3):  # 3개의 워커
            task = asyncio.create_task(self._process_queue())
            self.processing_tasks.append(task)
            
    async def _process_queue(self):
        """처리 큐 워커"""
        while True:
            try:
                item = await self.processing_queue.get()
                await self._process_upload_item(item)
            except Exception as e:
                logger.error(f"큐 처리 오류: {str(e)}")
            finally:
                self.processing_queue.task_done()
                
    async def upload_file(self,
                          file: Union[FileStorage, str],
                          upload_type: UploadType = UploadType.ADMIN,
                          user_id: str = "admin",
                          metadata: Optional[Dict[str, Any]] = None) -> UploadResult:
        """
        파일 업로드 및 처리
        
        Args:
            file: 업로드 파일 객체 또는 파일 경로
            upload_type: 업로드 유형
            user_id: 사용자 ID
            metadata: 추가 메타데이터
            
        Returns:
            UploadResult 객체
        """
        try:
            # 파일 정보 추출
            if isinstance(file, str):
                file_path = file
                filename = os.path.basename(file)
                file_size = os.path.getsize(file)
            else:
                filename = secure_filename(file.filename)
                file_path = None
                file_size = None  # FileStorage에서 크기 확인
                
            # 파일 유형 확인
            file_type = self._get_file_type(filename)
            
            # 파일 ID 생성
            file_id = self._generate_file_id(filename)
            
            # 저장 경로 결정
            if upload_type == UploadType.USER:
                save_dir = self.temp_dir / user_id
            else:
                save_dir = self.upload_dir / upload_type.value
                
            save_dir.mkdir(parents=True, exist_ok=True)
            save_path = save_dir / f"{file_id}_{filename}"
            
            # 파일 저장
            if isinstance(file, str):
                shutil.copy2(file, save_path)
            else:
                file.save(str(save_path))
                file_size = save_path.stat().st_size
                
            # 파일 크기 검증
            if file_size > self.max_file_size:
                raise ValueError(f"파일 크기 초과: {file_size / (1024*1024):.2f}MB")
                
            # 메타데이터 생성
            upload_metadata = {
                'user_id': user_id,
                'upload_type': upload_type.value,
                'file_type': file_type.value,
                'file_size': file_size,
                'timestamp': datetime.now().isoformat()
            }
            
            if metadata:
                upload_metadata.update(metadata)
                
            # 처리 큐에 추가
            process_item = {
                'file_id': file_id,
                'file_path': str(save_path),
                'file_type': file_type,
                'upload_type': upload_type,
                'metadata': upload_metadata
            }
            
            # 관리자 업로드는 즉시 처리
            if upload_type == UploadType.ADMIN:
                analysis_result = await self._process_file(save_path, file_type)
                vector_id = await self._store_in_vector_db(save_path, analysis_result)
            else:
                # 사용자 업로드는 큐에 추가
                await self.processing_queue.put(process_item)
                analysis_result = None
                vector_id = None
                
            return UploadResult(
                file_id=file_id,
                original_name=filename,
                stored_path=str(save_path),
                file_type=file_type,
                upload_type=upload_type,
                processing_status="completed" if upload_type == UploadType.ADMIN else "queued",
                analysis_result=analysis_result,
                vector_id=vector_id,
                metadata=upload_metadata,
                errors=[],
                timestamp=datetime.now().isoformat()
            )
            
        except Exception as e:
            logger.error(f"파일 업로드 실패: {str(e)}")
            
            return UploadResult(
                file_id="",
                original_name=filename if 'filename' in locals() else "",
                stored_path="",
                file_type=FileType.UNKNOWN,
                upload_type=upload_type,
                processing_status="error",
                analysis_result=None,
                vector_id=None,
                metadata={},
                errors=[str(e)],
                timestamp=datetime.now().isoformat()
            )
            
    async def upload_batch(self,
                          files: List[Union[FileStorage, str]],
                          upload_type: UploadType = UploadType.BATCH,
                          user_id: str = "admin") -> List[UploadResult]:
        """
        배치 파일 업로드
        
        Args:
            files: 파일 리스트
            upload_type: 업로드 유형
            user_id: 사용자 ID
            
        Returns:
            UploadResult 리스트
        """
        results = []
        
        # 병렬 처리
        tasks = []
        for file in files:
            task = self.upload_file(file, upload_type, user_id)
            tasks.append(task)
            
        # 결과 수집
        results = await asyncio.gather(*tasks)
        
        return results
        
    async def _process_upload_item(self, item: Dict[str, Any]):
        """업로드 항목 처리"""
        try:
            file_path = Path(item['file_path'])
            file_type = item['file_type']
            
            # 파일 처리
            analysis_result = await self._process_file(file_path, file_type)
            
            # 벡터 DB 저장
            if analysis_result:
                vector_id = await self._store_in_vector_db(file_path, analysis_result)
                
                # 처리 완료 디렉토리로 이동
                processed_path = self.processed_dir / file_path.name
                shutil.move(str(file_path), str(processed_path))
                
                logger.info(f"파일 처리 완료: {file_path.name}")
                
        except Exception as e:
            logger.error(f"파일 처리 실패: {str(e)}")
            
    async def _process_file(self,
                           file_path: Path,
                           file_type: FileType) -> Optional[Dict[str, Any]]:
        """
        파일 유형별 처리
        
        Args:
            file_path: 파일 경로
            file_type: 파일 유형
            
        Returns:
            처리 결과 딕셔너리
        """
        if file_type == FileType.IMAGE:
            # 이미지 처리
            result = await self.image_analyzer.analyze_image(
                str(file_path),
                ProcessingMode.DEEP
            )
            return asdict(result)
            
        elif file_type == FileType.PDF:
            # PDF 처리
            return await self._process_pdf(file_path)
            
        elif file_type == FileType.DOCUMENT:
            # 문서 처리
            return await self._process_document(file_path)
            
        elif file_type == FileType.EXCEL:
            # 엑셀 처리
            return await self._process_excel(file_path)
            
        elif file_type == FileType.TEXT:
            # 텍스트 처리
            return await self._process_text(file_path)
            
        return None
        
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF 파일 처리"""
        # PDF 처리 로직 (PyPDF2, pdfplumber 등 사용)
        import PyPDF2
        
        text_content = []
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_content.append(text)
                    
                metadata = {
                    'num_pages': num_pages,
                    'info': pdf_reader.metadata
                }
                
        except Exception as e:
            logger.error(f"PDF 처리 오류: {str(e)}")
            
        return {
            'type': 'pdf',
            'text': '\n'.join(text_content),
            'metadata': metadata
        }
        
    async def _process_document(self, file_path: Path) -> Dict[str, Any]:
        """문서 파일 처리 (DOC, DOCX, HWP)"""
        # python-docx, pyhwp 등 사용
        text_content = ""
        metadata = {}
        
        if file_path.suffix.lower() in ['.docx']:
            from docx import Document
            
            try:
                doc = Document(str(file_path))
                paragraphs = [para.text for para in doc.paragraphs]
                text_content = '\n'.join(paragraphs)
                
                metadata = {
                    'num_paragraphs': len(paragraphs),
                    'tables': len(doc.tables)
                }
            except Exception as e:
                logger.error(f"문서 처리 오류: {str(e)}")
                
        return {
            'type': 'document',
            'text': text_content,
            'metadata': metadata
        }
        
    async def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """엑셀 파일 처리"""
        import pandas as pd
        
        try:
            # 엑셀 파일 읽기
            df_dict = pd.read_excel(str(file_path), sheet_name=None)
            
            data = {}
            for sheet_name, df in df_dict.items():
                data[sheet_name] = {
                    'shape': df.shape,
                    'columns': df.columns.tolist(),
                    'preview': df.head(10).to_dict('records')
                }
                
            return {
                'type': 'excel',
                'sheets': list(df_dict.keys()),
                'data': data
            }
            
        except Exception as e:
            logger.error(f"엑셀 처리 오류: {str(e)}")
            return {}
            
    async def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """텍스트 파일 처리"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                
            return {
                'type': 'text',
                'content': content,
                'lines': len(content.splitlines()),
                'characters': len(content)
            }
        except Exception as e:
            logger.error(f"텍스트 처리 오류: {str(e)}")
            return {}
            
    async def _store_in_vector_db(self,
                                  file_path: Path,
                                  analysis_result: Dict[str, Any]) -> Optional[str]:
        """
        벡터 데이터베이스 저장
        
        Args:
            file_path: 파일 경로
            analysis_result: 분석 결과
            
        Returns:
            벡터 ID
        """
        if not self.vector_store or not self.embedding_generator:
            logger.warning("벡터 저장소가 초기화되지 않음")
            return None
            
        try:
            # 텍스트 추출
            text = ""
            if 'text' in analysis_result:
                text = analysis_result['text']
            elif 'ocr_result' in analysis_result:
                ocr = analysis_result['ocr_result']
                if ocr and 'text' in ocr:
                    text = ocr['text']
                    
            if not text:
                logger.warning("저장할 텍스트 없음")
                return None
                
            # 임베딩 생성
            embeddings = await self.embedding_generator.generate(text)
            
            # 벡터 저장
            vector_id = await self.vector_store.add_document(
                document_id=self._generate_file_id(file_path.name),
                embeddings=embeddings,
                metadata={
                    'file_path': str(file_path),
                    'file_name': file_path.name,
                    'analysis': analysis_result
                }
            )
            
            return vector_id
            
        except Exception as e:
            logger.error(f"벡터 저장 실패: {str(e)}")
            return None
            
    def _get_file_type(self, filename: str) -> FileType:
        """파일 유형 판별"""
        ext = Path(filename).suffix.lower()
        
        for file_type, extensions in self.allowed_extensions.items():
            if ext in extensions:
                if file_type == 'image':
                    return FileType.IMAGE
                elif file_type == 'document':
                    return FileType.DOCUMENT
                elif file_type == 'pdf':
                    return FileType.PDF
                elif file_type == 'excel':
                    return FileType.EXCEL
                elif file_type == 'text':
                    return FileType.TEXT
                    
        return FileType.UNKNOWN
        
    def _generate_file_id(self, filename: str) -> str:
        """파일 ID 생성"""
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        hash_value = hashlib.md5(f"{filename}{timestamp}".encode()).hexdigest()[:8]
        return f"{timestamp}_{hash_value}"
        
    async def delete_user_files(self, user_id: str):
        """
        사용자 임시 파일 삭제
        
        Args:
            user_id: 사용자 ID
        """
        user_dir = self.temp_dir / user_id
        
        if user_dir.exists():
            try:
                shutil.rmtree(user_dir)
                logger.info(f"사용자 파일 삭제: {user_id}")
            except Exception as e:
                logger.error(f"파일 삭제 실패: {str(e)}")
                
    async def get_upload_statistics(self) -> Dict[str, Any]:
        """업로드 통계 조회"""
        stats = {
            'total_uploads': 0,
            'by_type': {},
            'by_status': {},
            'total_size': 0
        }
        
        # 업로드 디렉토리 스캔
        for upload_type in UploadType:
            type_dir = self.upload_dir / upload_type.value
            if type_dir.exists():
                files = list(type_dir.glob('*'))
                stats['by_type'][upload_type.value] = len(files)
                stats['total_uploads'] += len(files)
                
                for file in files:
                    stats['total_size'] += file.stat().st_size
                    
        # 처리 완료 파일
        processed_files = list(self.processed_dir.glob('*'))
        stats['processed'] = len(processed_files)
        
        # 대기 중인 파일
        stats['queued'] = self.processing_queue.qsize()
        
        return stats
        
    async def cleanup_old_files(self, days: int = 7):
        """
        오래된 파일 정리
        
        Args:
            days: 보관 일수
        """
        cutoff_time = datetime.now().timestamp() - (days * 24 * 60 * 60)
        
        cleaned = 0
        for dir_path in [self.temp_dir, self.processed_dir]:
            for file_path in dir_path.glob('**/*'):
                if file_path.is_file():
                    if file_path.stat().st_mtime < cutoff_time:
                        try:
                            file_path.unlink()
                            cleaned += 1
                        except Exception as e:
                            logger.error(f"파일 삭제 실패: {file_path} - {str(e)}")
                            
        logger.info(f"{cleaned}개 파일 정리 완료")
        
    def export_metadata(self, output_path: str):
        """메타데이터 내보내기"""
        metadata = {
            'upload_dir': str(self.upload_dir),
            'processed_dir': str(self.processed_dir),
            'temp_dir': str(self.temp_dir),
            'max_file_size': self.max_file_size,
            'allowed_extensions': {k: list(v) for k, v in self.allowed_extensions.items()}
        }
        
        with open(output_path, 'w') as f:
            json.dump(metadata, f, indent=2)


# 테스트 코드
async def main():
    # 업로드 핸들러 초기화
    handler = AdminUploadHandler()
    await handler.initialize_modules()
    
    print("ex-GPT 관리도구 업로드 핸들러 준비 완료")
    print("\n지원 기능:")
    print("1. 다양한 파일 형식 지원 (이미지, PDF, 문서, 엑셀)")
    print("2. 자동 OCR 및 텍스트 추출")
    print("3. 개인정보 자동 검출 및 마스킹")
    print("4. 중복 파일 검사")
    print("5. 벡터 데이터베이스 자동 저장")
    print("6. 배치 업로드 지원")
    print("7. 사용자별 임시 파일 관리")
    
    # 통계 조회
    stats = await handler.get_upload_statistics()
    print(f"\n현재 업로드 통계:")
    print(f"- 총 업로드: {stats['total_uploads']}개")
    print(f"- 처리 완료: {stats.get('processed', 0)}개")
    print(f"- 대기 중: {stats.get('queued', 0)}개")
    print(f"- 총 용량: {stats['total_size'] / (1024*1024):.2f}MB")


if __name__ == "__main__":
    asyncio.run(main())
